"""
NOVA Tool System (Phase 3)
===========================
Enterprise-grade tool execution with security
"""

import subprocess
import psutil
import os
from pathlib import Path
from loguru import logger
from typing import Dict, Any, Optional, List
import json


class ToolExecutor:
    """Enterprise tool executor with permission system"""
    
    def __init__(self, config, security_manager):
        self.config = config
        self.security = security_manager
        self.tools = self._register_tools()
        logger.info(f"Tool executor initialized with {len(self.tools)} tools")
    
    def _register_tools(self) -> Dict[str, callable]:
        """Register all available tools"""
        return {
            # System Control
            "open_app": self.open_app,
            "close_app": self.close_app,
            "list_apps": self.list_running_apps,
            
            # File Operations
            "search_file": self.search_file,
            "create_file": self.create_file,
            "read_file": self.read_file,
            "delete_file": self.delete_file,
            "move_file": self.move_file,
            "create_folder": self.create_folder,
            
            # Document Processing
            "read_pdf": self.read_pdf,
            "summarize_pdf": self.summarize_pdf,
            "create_document": self.create_document,
            
            # Command Execution
            "run_command": self.run_command,
            "get_system_info": self.get_system_info,
            
            # Web & Search
            "open_url": self.open_url,
            "search_web": self.search_web,
        }
    
    def execute(self, tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """
        Execute a tool with security checks
        
        Returns:
            {
                "success": bool,
                "result": any,
                "message": str
            }
        """
        # Check if tool exists
        if tool_name not in self.tools:
            return {
                "success": False,
                "result": None,
                "message": f"Unknown tool: {tool_name}"
            }
        
        # Security check
        if not self.security.can_execute(tool_name, parameters):
            return {
                "success": False,
                "result": None,
                "message": f"Permission denied for {tool_name}"
            }
        
        # Execute tool
        try:
            logger.info(f"Executing tool: {tool_name} with {parameters}")
            result = self.tools[tool_name](**parameters)
            
            # Log successful execution
            self.security.log_execution(tool_name, parameters, success=True)
            
            return {
                "success": True,
                "result": result,
                "message": "Success"
            }
        
        except Exception as e:
            error_msg = f"Tool execution failed: {e}"
            logger.error(error_msg)
            self.security.log_execution(tool_name, parameters, success=False, error=str(e))
            
            return {
                "success": False,
                "result": None,
                "message": error_msg
            }
    
    # ==================== SYSTEM CONTROL TOOLS ====================
    
    def open_app(self, app_name: str) -> str:
        """Open an application"""
        app_map = {
            'chrome': 'chrome.exe',
            'firefox': 'firefox.exe',
            'edge': 'msedge.exe',
            'notepad': 'notepad.exe',
            'vscode': 'code.exe',
            'explorer': 'explorer.exe',
            'calc': 'calc.exe',
            'cmd': 'cmd.exe',
            'powershell': 'powershell.exe',
        }
        
        app_exe = app_map.get(app_name.lower(), app_name)
        
        try:
            subprocess.Popen([app_exe])
            return f"Opened {app_name}"
        except Exception as e:
            raise Exception(f"Could not open {app_name}: {e}")
    
    def close_app(self, app_name: str) -> str:
        """Close an application"""
        closed_count = 0
        
        for proc in psutil.process_iter(['name']):
            try:
                if app_name.lower() in proc.info['name'].lower():
                    proc.terminate()
                    closed_count += 1
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        
        if closed_count > 0:
            return f"Closed {closed_count} instance(s) of {app_name}"
        else:
            return f"No running instances of {app_name} found"
    
    def list_running_apps(self) -> List[str]:
        """List all running applications"""
        apps = []
        for proc in psutil.process_iter(['name']):
            try:
                apps.append(proc.info['name'])
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        return list(set(apps))[:20]  # Return unique apps, max 20
    
    # ==================== FILE OPERATIONS ====================
    
    def search_file(self, query: str, directory: str = None) -> List[str]:
        """Search for files matching query"""
        import glob
        
        if directory is None:
            directory = str(Path.home())
        
        search_path = Path(directory)
        results = []
        
        # Search in common locations
        patterns = [
            f"**/*{query}*",
            f"**/{query}*",
        ]
        
        for pattern in patterns:
            for file in search_path.glob(pattern):
                if file.is_file():
                    results.append(str(file))
                if len(results) >= 10:  # Limit results
                    break
            if results:
                break
        
        return results
    
    def create_file(self, path: str, content: str = "") -> str:
        """Create a new file"""
        file_path = Path(path)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text(content, encoding='utf-8')
        return f"Created file: {path}"
    
    def read_file(self, path: str, max_lines: int = 100) -> str:
        """Read file contents"""
        file_path = Path(path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        content = file_path.read_text(encoding='utf-8')
        lines = content.split('\n')
        
        if len(lines) > max_lines:
            return '\n'.join(lines[:max_lines]) + f"\n... ({len(lines) - max_lines} more lines)"
        return content
    
    def delete_file(self, path: str) -> str:
        """Delete a file"""
        file_path = Path(path)
        
        if file_path.exists():
            file_path.unlink()
            return f"Deleted: {path}"
        else:
            return f"File not found: {path}"
    
    def move_file(self, source: str, destination: str) -> str:
        """Move a file"""
        import shutil
        shutil.move(source, destination)
        return f"Moved {source} to {destination}"
    
    def create_folder(self, path: str) -> str:
        """Create a directory"""
        Path(path).mkdir(parents=True, exist_ok=True)
        return f"Created folder: {path}"
    
    # ==================== DOCUMENT PROCESSING ====================
    
    def read_pdf(self, path: str, max_pages: int = 5) -> str:
        """Read PDF content"""
        try:
            import PyPDF2
            
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                content = []
                for i in range(min(num_pages, max_pages)):
                    page = pdf_reader.pages[i]
                    content.append(page.extract_text())
                
                result = '\n\n'.join(content)
                if num_pages > max_pages:
                    result += f"\n\n... (Total {num_pages} pages, showing first {max_pages})"
                
                return result
        
        except Exception as e:
            raise Exception(f"Could not read PDF: {e}")
    
    def summarize_pdf(self, path: str, llm=None) -> str:
        """Summarize PDF using LLM"""
        content = self.read_pdf(path)
        
        if llm:
            summary = llm.get_quick_response(
                f"Summarize this PDF content in 3-5 sentences:\n\n{content[:2000]}"
            )
            return summary
        else:
            return f"PDF preview (no LLM for summary): {content[:500]}..."
    
    def create_document(self, path: str, title: str, content: str, doc_type: str = "txt") -> str:
        """Create a document (txt, docx, etc)"""
        if doc_type == "docx":
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(content)
                doc.save(path)
                return f"Created Word document: {path}"
            except ImportError:
                raise Exception("python-docx not installed")
        else:
            # TXT file
            self.create_file(path, f"{title}\n\n{content}")
            return f"Created text document: {path}"
    
    # ==================== COMMAND EXECUTION ====================
    
    def run_command(self, command: str, timeout: int = 30) -> str:
        """Execute a shell command (with restrictions)"""
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=timeout
        )
        
        output = result.stdout if result.stdout else result.stderr
        return output[:1000]  # Limit output
    
    def get_system_info(self) -> Dict[str, Any]:
        """Get system information"""
        import platform
        
        cpu_percent = psutil.cpu_percent(interval=1)
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        return {
            "os": platform.system(),
            "os_version": platform.version(),
            "cpu_percent": cpu_percent,
            "memory_percent": memory.percent,
            "memory_available_gb": memory.available / (1024**3),
            "disk_percent": disk.percent,
            "disk_free_gb": disk.free / (1024**3),
        }
    
    # ==================== WEB & SEARCH ====================
    
    def open_url(self, url: str) -> str:
        """Open URL in default browser"""
        import webbrowser
        webbrowser.open(url)
        return f"Opened URL: {url}"
    
    def search_web(self, query: str) -> str:
        """Open web search"""
        import webbrowser
        search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
        webbrowser.open(search_url)
        return f"Opened search for: {query}"


if __name__ == "__main__":
    # Test tools
    from config.config import config
    from tools.security import SecurityManager
    
    security = SecurityManager(config)
    executor = ToolExecutor(config, security)
    
    # Test file creation
    result = executor.execute("create_file", {
        "path": "test_output.txt",
        "content": "Hello from NOVA tool system!"
    })
    print(result)
    
    # Test system info
    result = executor.execute("get_system_info", {})
    print(result)
